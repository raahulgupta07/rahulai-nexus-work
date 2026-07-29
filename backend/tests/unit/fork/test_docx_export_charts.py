"""Word export must carry a doc's embedded charts as real pictures.

The bug this locks: `markdown_to_docx` deleted every ``{{viz:<uuid>}}``
placeholder with a regex, so an exported .docx of a document containing three
live charts came back valid, ~43KB, and with `inline_shapes == 0` — no chart,
no table, no note that anything was dropped.

Fast + self-contained: no DB, no network, no browser. The headless ECharts
render is exercised at the seam (assets in), which is exactly where the old
code threw the charts away.
"""
import io
import struct
import zlib

import pytest
from docx import Document

from app.services.docx_export_service import markdown_to_docx


VIZ_A = "11111111-2222-3333-4444-555555555555"
VIZ_B = "aaaaaaaa-bbbb-cccc-dddd-eeeeeeeeeeee"


def _png_bytes(width: int = 40, height: int = 30) -> bytes:
    """A minimal, valid PNG — python-docx reads the header to size the picture."""
    def chunk(tag: bytes, data: bytes) -> bytes:
        return (
            struct.pack(">I", len(data))
            + tag
            + data
            + struct.pack(">I", zlib.crc32(tag + data) & 0xFFFFFFFF)
        )

    ihdr = struct.pack(">IIBBBBB", width, height, 8, 2, 0, 0, 0)
    raw = b"".join(b"\x00" + b"\x7f\x7f\x7f" * width for _ in range(height))
    return (
        b"\x89PNG\r\n\x1a\n"
        + chunk(b"IHDR", ihdr)
        + chunk(b"IDAT", zlib.compress(raw))
        + chunk(b"IEND", b"")
    )


def _doc(markdown: str, viz_assets=None) -> Document:
    return Document(io.BytesIO(markdown_to_docx(markdown, "Report", viz_assets=viz_assets)))


def _text(doc) -> str:
    return "\n".join(p.text for p in doc.paragraphs)


THREE_CHART_DOC = f"""# Quarterly review

Revenue climbed through the quarter.

{{{{viz:{VIZ_A}}}}}

Costs held flat.

{{{{viz:{VIZ_B}}}}}

{{{{viz:{VIZ_A}}}}}
"""


def test_three_embedded_charts_become_three_pictures():
    """The reported bug: three embedded charts, zero inline shapes."""
    png = _png_bytes()
    assets = {
        VIZ_A: {"title": "Revenue by month", "image_png": png, "table": None},
        VIZ_B: {"title": "Cost by month", "image_png": png, "table": None},
    }
    doc = _doc(THREE_CHART_DOC, assets)

    assert len(doc.inline_shapes) == 3


def test_chart_titles_render_as_captions():
    png = _png_bytes()
    assets = {VIZ_A: {"title": "Revenue by month", "image_png": png, "table": None}}
    doc = _doc(f"Intro\n\n{{{{viz:{VIZ_A}}}}}\n", assets)

    assert len(doc.inline_shapes) == 1
    assert "Revenue by month" in _text(doc)


def test_placeholder_is_never_left_as_raw_text():
    """No assets at all: the reader gets a note, not `{{viz:...}}` or a stray `{}`."""
    doc = _doc(f"Intro\n\n{{{{viz:{VIZ_A}}}}}\n")

    body = _text(doc)
    assert "viz:" not in body
    assert "{" not in body and "}" not in body
    assert "[chart unavailable]" in body


def test_table_visualization_becomes_a_word_table():
    assets = {
        VIZ_A: {
            "title": "Top branches",
            "image_png": None,
            "table": {
                "columns": [
                    {"field": "branch", "headerName": "Branch"},
                    {"field": "sales", "headerName": "Sales"},
                ],
                "rows": [
                    {"branch": "Yangon", "sales": 120},
                    {"branch": "Mandalay", "sales": 90},
                ],
            },
        }
    }
    doc = _doc(f"Intro\n\n{{{{viz:{VIZ_A}}}}}\n", assets)

    assert len(doc.tables) == 1
    table = doc.tables[0]
    assert table.cell(0, 0).text == "Branch"
    assert table.cell(1, 0).text == "Yangon"
    assert table.cell(2, 1).text == "90"


def test_table_columns_fall_back_to_row_keys():
    assets = {
        VIZ_A: {
            "title": "",
            "image_png": None,
            "table": {"columns": [], "rows": [{"city": "Yangon", "n": 3}]},
        }
    }
    doc = _doc(f"{{{{viz:{VIZ_A}}}}}\n", assets)

    headers = {doc.tables[0].cell(0, c).text for c in range(2)}
    assert headers == {"city", "n"}


def test_wide_raw_dataset_is_capped_not_dumped():
    """The regression this replaced: 51 rows x 16 cols wrapped to one char per line.

    A table embed is a preview in a printed document. Cap columns AND rows, and
    say so — never dump the dataset.
    """
    columns = [{"field": f"c{i}", "headerName": f"Column {i}"} for i in range(16)]
    rows = [{f"c{i}": f"v{r}_{i}" for i in range(16)} for r in range(51)]
    assets = {VIZ_A: {"title": "Board summary", "image_png": None,
                      "table": {"columns": columns, "rows": rows}}}
    doc = _doc(f"{{{{viz:{VIZ_A}}}}}\n", assets)

    table = doc.tables[0]
    assert len(table.columns) <= 6
    assert len(table.rows) <= 13  # header + 12 body rows
    caption = _text(doc)
    assert "12 of 51 rows" in caption and "6 of 16 columns" in caption


def test_long_cell_values_are_clipped():
    assets = {
        VIZ_A: {
            "title": "",
            "image_png": None,
            "table": {"columns": [{"field": "name", "headerName": "Name"}],
                      "rows": [{"name": "x" * 200}]},
        }
    }
    doc = _doc(f"{{{{viz:{VIZ_A}}}}}\n", assets)

    assert len(doc.tables[0].cell(1, 0).text) <= 40


def test_small_table_is_not_captioned_as_truncated():
    assets = {
        VIZ_A: {
            "title": "",
            "image_png": None,
            "table": {"columns": [{"field": "a", "headerName": "A"}], "rows": [{"a": 1}]},
        }
    }
    doc = _doc(f"{{{{viz:{VIZ_A}}}}}\n", assets)

    assert "Preview" not in _text(doc)


def test_broken_image_falls_back_to_the_table_free_note():
    """An unreadable PNG must not abort the export."""
    assets = {VIZ_A: {"title": "Revenue", "image_png": b"not-a-png", "table": None}}
    doc = _doc(f"Body\n\n{{{{viz:{VIZ_A}}}}}\n", assets)

    assert len(doc.inline_shapes) == 0
    assert "[chart unavailable: Revenue]" in _text(doc)
    assert "Body" in _text(doc)


def test_inline_placeholder_keeps_its_paragraph_text():
    png = _png_bytes()
    assets = {VIZ_A: {"title": "", "image_png": png, "table": None}}
    doc = _doc(f"See the chart {{{{viz:{VIZ_A}}}}} above.\n", assets)

    assert len(doc.inline_shapes) == 1
    assert "See the chart" in _text(doc)
    assert "above." in _text(doc)


def test_quoted_placeholder_in_a_code_fence_is_not_an_embed():
    png = _png_bytes()
    assets = {VIZ_A: {"title": "", "image_png": png, "table": None}}
    markdown = f"How to embed:\n\n```\n{{{{viz:{VIZ_A}}}}}\n```\n"
    doc = _doc(markdown, assets)

    assert len(doc.inline_shapes) == 0
    assert f"viz:{VIZ_A}" in _text(doc)


def test_ordinary_markdown_still_converts():
    doc = _doc(
        "# Title\n\n## Section\n\n- one\n- two\n\n1. first\n\n"
        "| A | B |\n| --- | --- |\n| 1 | 2 |\n\nPlain **bold** text.\n"
    )

    body = _text(doc)
    assert "Title" in body and "Section" in body
    assert "one" in body and "first" in body
    assert len(doc.tables) == 1
    assert doc.tables[0].cell(0, 0).text == "A"


@pytest.mark.parametrize("markdown", ["", None])
def test_empty_document_does_not_raise(markdown):
    assert markdown_to_docx(markdown, "Report").startswith(b"PK")
