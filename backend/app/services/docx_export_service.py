"""Markdown → .docx export for `doc` artifacts.

Minimal, dependency-light converter covering what the doc-writer actually
emits: #/##/### headings, paragraphs, - / * bullets, 1. numbered lists,
**bold** / *italic* / `code` inline, GitHub-style | pipe | tables, and
``{{viz:<uuid>}}`` chart embeds.
Anything unrecognized degrades to a plain paragraph — never raises.

Chart embeds: the caller (the export route) resolves each ``{{viz:<uuid>}}``
placeholder to an asset via `app.services.doc_viz_render` and passes them in
as `viz_assets`. A chart asset carries PNG bytes (rendered headlessly with the
same ECharts build the app uses) and becomes a real inline picture; a
table/metric asset becomes a native Word table. Without assets the placeholder
degrades to a one-line note — it is never left as raw `{{viz:...}}` text.
"""
import io
import re
from typing import Any, Dict, List, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


_INLINE_RE = re.compile(r"(\*\*.+?\*\*|\*.+?\*|`.+?`)")

# Canonical embed syntax emitted by create_doc/edit_doc and rendered by
# DocViewer: {{viz:<uuid>}}. Kept identical to
# `app.ai.tools.implementations._doc_markdown._VIZ_PLACEHOLDER_RE`.
_VIZ_EMBED_RE = re.compile(r"\{\{\s*viz:\s*([0-9a-fA-F-]{8,64})\s*\}\}")

# Body text is 12px-ish in the viewer; 6" keeps a chart inside A4/Letter margins.
_PICTURE_WIDTH = Inches(6.0)

# A table embed is a PREVIEW in a printed document, not the dataset. A 16-column
# table on A4 wraps every cell to one character per line — unreadable, and it
# turned a 4-page report into 18. Cap hard and say so in a caption; the full data
# lives in the report.
_MAX_TABLE_ROWS = 12
_MAX_TABLE_COLS = 6
_MAX_CELL_CHARS = 40
_TABLE_FONT_SIZE = Pt(8)


def _add_inline(par, text: str) -> None:
    """Render **bold** / *italic* / `code` runs inside a paragraph."""
    for part in _INLINE_RE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**") and len(part) > 4:
            par.add_run(part[2:-2]).bold = True
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            par.add_run(part[1:-1]).italic = True
        elif part.startswith("`") and part.endswith("`") and len(part) > 2:
            run = par.add_run(part[1:-1])
            run.font.name = "Courier New"
            run.font.size = Pt(9)
        else:
            par.add_run(part)


def _is_table_row(line: str) -> bool:
    s = line.strip()
    return s.startswith("|") and s.endswith("|") and s.count("|") >= 2


def _is_divider_row(line: str) -> bool:
    return bool(re.fullmatch(r"\|(\s*:?-+:?\s*\|)+", line.strip()))


def _clip(text: str) -> str:
    """Keep one cell from swallowing a column width."""
    text = (text or "").strip()
    return text if len(text) <= _MAX_CELL_CHARS else text[: _MAX_CELL_CHARS - 1] + "…"


def _table_font_for(ncols: int):
    """Body size for a narrow table; step down as columns are added.

    Returns None for tables narrow enough to leave at the document's own size —
    a two-column table shrunk to 8pt looks broken for no reason.
    """
    if ncols <= 4:
        return None
    if ncols <= 6:
        return Pt(9)
    if ncols <= 8:
        return Pt(8)
    return Pt(7)


def _no_wrap(cell) -> None:
    """Ask Word to widen the column rather than wrap the cell.

    `w:noWrap` is the .docx counterpart of the `white-space: nowrap` the PDF
    export uses, and it is here for the same reason: a wrapped data cell is not
    just ugly, it splits the value itself. Word honours this while the table
    still fits the page and falls back to wrapping when it cannot, so it can
    tighten a table but never push one off the edge. Header cells are left
    wrappable on purpose — a long column NAME may wrap without harm, and letting
    it force the column wide is what squeezes the data columns.
    """
    try:
        tc_pr = cell._tc.get_or_add_tcPr()
        if tc_pr.find(qn("w:noWrap")) is None:
            tc_pr.append(OxmlElement("w:noWrap"))
    except Exception:
        # Cosmetic only — never lose a document over a table hint.
        pass


def _add_caption(doc, text: str) -> None:
    par = doc.add_paragraph()
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = par.add_run(text)
    run.italic = True
    run.font.size = Pt(9)


def _add_data_table(doc, table_data: Dict[str, Any]) -> bool:
    """Render a viz's rows as a native Word table. True when anything was added."""
    rows = (table_data or {}).get("rows") or []
    columns = (table_data or {}).get("columns") or []
    if not rows:
        return False

    # Column descriptors follow the app's grid shape ({headerName, field}); fall
    # back to the first row's keys when a viz carries rows but no column spec.
    fields: List[str] = []
    headers: List[str] = []
    for col in columns:
        if isinstance(col, dict):
            field = col.get("field") or col.get("headerName")
            if not field:
                continue
            fields.append(str(field))
            headers.append(str(col.get("headerName") or field))
        elif col:
            fields.append(str(col))
            headers.append(str(col))
    if not fields:
        first = rows[0]
        if not isinstance(first, dict):
            return False
        fields = [str(k) for k in first.keys()]
        headers = list(fields)

    total_cols = len(fields)
    total_rows = len(rows)
    body_rows = rows[:_MAX_TABLE_ROWS]

    # Spend the few columns a page can hold on columns that carry data. The
    # reported dump wasted 2 of 6 on fields that are blank for every previewed
    # row (banner/outlet_name on quarter-total rows).
    def _has_data(field: str) -> bool:
        for row in body_rows:
            value = row.get(field) if isinstance(row, dict) else None
            if value is not None and str(value).strip() != "":
                return True
        return False

    kept = [i for i, f in enumerate(fields) if _has_data(f)] or list(range(total_cols))
    shown_fields = [fields[i] for i in kept][:_MAX_TABLE_COLS]
    shown_headers = [headers[i] for i in kept][:_MAX_TABLE_COLS]

    table = doc.add_table(rows=len(body_rows) + 1, cols=len(shown_fields))
    table.style = "Light Grid Accent 1"
    for c_idx, header in enumerate(shown_headers):
        run = table.cell(0, c_idx).paragraphs[0].add_run(_clip(header))
        run.bold = True
        run.font.size = _TABLE_FONT_SIZE
    for r_idx, row in enumerate(body_rows, start=1):
        for c_idx, field in enumerate(shown_fields):
            value = row.get(field) if isinstance(row, dict) else None
            run = table.cell(r_idx, c_idx).paragraphs[0].add_run(
                "" if value is None else _clip(str(value))
            )
            run.font.size = _TABLE_FONT_SIZE

    # Never let a reader mistake a preview for the whole thing.
    notes = []
    if total_rows > len(body_rows):
        notes.append(f"{len(body_rows)} of {total_rows} rows")
    if total_cols > len(shown_fields):
        notes.append(f"{len(shown_fields)} of {total_cols} columns")
    if notes:
        _add_caption(doc, "Preview — showing " + " and ".join(notes) + ". Full data in the report.")
    return True


def _add_viz(doc, viz_id: str, viz_assets: Optional[Dict[str, Any]]) -> None:
    """Render one ``{{viz:<uuid>}}`` embed into the document."""
    asset = (viz_assets or {}).get((viz_id or "").lower()) or {}
    viz_title = asset.get("title") or ""
    image_png = asset.get("image_png")

    if image_png:
        try:
            doc.add_picture(io.BytesIO(image_png), width=_PICTURE_WIDTH)
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            if viz_title:
                _add_caption(doc, viz_title)
            return
        except Exception:
            # A single unreadable image must not lose the rest of the document.
            pass

    if _add_data_table(doc, asset.get("table")):
        if viz_title:
            _add_caption(doc, viz_title)
        return

    par = doc.add_paragraph()
    run = par.add_run(
        f"[chart unavailable: {viz_title}]" if viz_title else "[chart unavailable]"
    )
    run.italic = True


def markdown_to_docx(
    markdown_text: str,
    title: str,
    viz_assets: Optional[Dict[str, Any]] = None,
) -> bytes:
    """Convert doc markdown to .docx bytes.

    `viz_assets` maps a lower-cased viz uuid to
    ``{"title": str, "image_png": bytes | None, "table": {"columns", "rows"} | None}``.
    """
    doc = Document()
    doc.add_heading(title or "Document", level=0)

    lines: List[str] = (markdown_text or "").splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        # fenced code block
        if stripped.startswith("```"):
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1  # closing fence
            par = doc.add_paragraph()
            run = par.add_run("\n".join(code_lines))
            run.font.name = "Courier New"
            run.font.size = Pt(9)
            continue

        # table block
        if _is_table_row(stripped):
            rows = []
            while i < len(lines) and _is_table_row(lines[i].strip()):
                if not _is_divider_row(lines[i]):
                    cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                    rows.append(cells)
                i += 1
            if rows:
                ncols = max(len(r) for r in rows)
                table = doc.add_table(rows=len(rows), cols=ncols)
                table.style = "Light Grid Accent 1"
                # A table the author wrote keeps every column — unlike a viz
                # embed, these are content, not a preview of a dataset. What it
                # gets instead is room: at body size a wide table wraps inside
                # each cell, and a wrapped cell does not merely look cramped, it
                # BREAKS THE VALUE — an 8-column table rendered "17,726,384" as
                # "17,726,3" over "84" and "Mandalay" as "Mandala" over "y".
                # Same defect the PDF export carries `white-space: nowrap` for.
                font_size = _table_font_for(ncols)
                for r_idx, r in enumerate(rows):
                    for c_idx in range(ncols):
                        cell = table.cell(r_idx, c_idx)
                        cell_par = cell.paragraphs[0]
                        _add_inline(cell_par, r[c_idx] if c_idx < len(r) else "")
                        for run in cell_par.runs:
                            if font_size is not None:
                                run.font.size = font_size
                            if r_idx == 0:
                                run.bold = True
                        if r_idx > 0:
                            _no_wrap(cell)
            continue

        # headings
        m = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if m:
            level = min(len(m.group(1)), 4)
            h = doc.add_heading("", level=level)
            _add_inline(h, m.group(2).strip())
            i += 1
            continue

        # bullets
        m = re.match(r"^[-*]\s+(.*)$", stripped)
        if m:
            par = doc.add_paragraph(style="List Bullet")
            _add_inline(par, m.group(1))
            i += 1
            continue

        # numbered list
        m = re.match(r"^\d+[.)]\s+(.*)$", stripped)
        if m:
            par = doc.add_paragraph(style="List Number")
            _add_inline(par, m.group(1))
            i += 1
            continue

        # horizontal rule → skip
        if re.fullmatch(r"[-*_]{3,}", stripped):
            i += 1
            continue

        # plain paragraph. {{viz:<uuid>}} embeds become pictures/tables — the
        # old code deleted them (charts silently vanished from every export).
        viz_ids = [m.group(1) for m in _VIZ_EMBED_RE.finditer(stripped)]
        text = _VIZ_EMBED_RE.sub("", stripped).strip()
        if text:
            par = doc.add_paragraph()
            _add_inline(par, text)
        for viz_id in viz_ids:
            _add_viz(doc, viz_id, viz_assets)
        i += 1

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
